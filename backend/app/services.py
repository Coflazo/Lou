from __future__ import annotations

import io
import hashlib
import json
import re
import secrets
import uuid
import zipfile
from contextvars import ContextVar, Token
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from threading import Lock
from typing import Any

import fitz
import httpx
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from fastapi.encoders import jsonable_encoder
from fastapi.responses import StreamingResponse
from openpyxl import Workbook
from pypdf import PdfReader

from .ai import draft_contract_from_notes, parse_command_with_openai
from .algorithm_config import algorithm_config
from .algorithms import (
    BayesianRiskScorer,
    ClauseMatcher,
    CompanyBrainMindMap,
    HMMSectionDetector,
    VoiceMatcher,
)
from .algorithms.risk_scoring import RiskObservation
from .config import settings
from .provider_keys import current_slng_key
from sqlmodel import Session, select

from .db import ApiKeyRecord, engine, init_db, write_snapshot
from .models import (
    Commit,
    Contract,
    ContractFinding,
    FindingStatus,
    Playbook,
    PlaybookPosition,
    Proposal,
    ProposalStatus,
    Role,
)
from .seeder import seed_all


ROLE_RANK = {Role.JUNIOR: 1, Role.SENIOR: 2, Role.ADMIN: 3}


def _slng_stt_path() -> str:
    return f"/v1/stt/{settings.SLNG_STT_MODEL}"


def _slng_tts_path() -> str:
    return f"/v1/bridges/unmute/tts/{settings.SLNG_TTS_MODEL}"


def _slng_stt_http_url() -> str:
    return f"{settings.SLNG_API_BASE_HTTP}{_slng_stt_path()}"


def _slng_tts_http_url() -> str:
    return f"{settings.SLNG_API_BASE_HTTP}{_slng_tts_path()}"


def _slng_stt_ws_url() -> str:
    return f"{settings.SLNG_API_BASE_WS}{_slng_stt_path()}"


def _slng_tts_ws_url() -> str:
    return f"{settings.SLNG_API_BASE_WS}{_slng_tts_path()}"


@dataclass
class Store:
    current_role: Role = Role.JUNIOR
    playbooks: dict[str, Playbook] = field(default_factory=dict)
    contracts: dict[str, Contract] = field(default_factory=dict)
    proposals: dict[str, Proposal] = field(default_factory=dict)
    commits: dict[str, Commit] = field(default_factory=dict)
    entities: list[dict] = field(default_factory=list)
    relations: list[dict] = field(default_factory=list)

    def reset(self) -> None:
        data = seed_all()
        self.current_role = Role.JUNIOR
        self.playbooks = data["playbooks"]
        self.contracts = data["contracts"]
        self.proposals = data["proposals"]
        self.commits = data["commits"]
        self.entities = data["entities"]
        self.relations = data["relations"]


class AlgorithmRegistry:
    """Singleton holding per-playbook fitted matchers and the shared brain map."""

    def __init__(self) -> None:
        self._lock = Lock()
        self._matchers: dict[str, ClauseMatcher] = {}
        self._brain = CompanyBrainMindMap()
        self._brain.cache_ttl_seconds = float(algorithm_config.brain.cache_ttl_seconds)

    def get_matcher(self, playbook: Playbook) -> ClauseMatcher:
        with self._lock:
            matcher = self._matchers.get(playbook.id)
            if matcher is None:
                matcher = ClauseMatcher(min_score=algorithm_config.clause_matching.min_score)
                matcher.fit(playbook.positions)
                self._matchers[playbook.id] = matcher
            return matcher

    def invalidate(self, playbook_id: str | None = None) -> None:
        with self._lock:
            if playbook_id is None:
                self._matchers.clear()
            else:
                self._matchers.pop(playbook_id, None)
        self._brain.invalidate()

    def voice_matcher(self, playbook: Playbook) -> VoiceMatcher:
        vm = algorithm_config.voice_matching
        return VoiceMatcher(
            clause_matcher=self.get_matcher(playbook),
            threshold=vm.threshold,
            jaro_weight=vm.jaro_weight,
            tfidf_weight=vm.tfidf_weight,
            edit_weight=vm.edit_weight,
        )

    def risk_scorer(self) -> BayesianRiskScorer:
        rs = algorithm_config.risk_scoring
        return BayesianRiskScorer(
            levels=list(rs.levels),
            prior_alpha=list(rs.prior_alpha),
        )

    def section_detector(self) -> HMMSectionDetector:
        hmm = algorithm_config.hmm_section_detector
        return HMMSectionDetector(
            init_inside=hmm.init_inside,
            init_begin=hmm.init_begin,
            transition_inside_to_inside=hmm.transition_inside_to_inside,
            transition_begin_to_begin=hmm.transition_begin_to_begin,
            w_begin=hmm.logistic.w_begin,
            b_begin=hmm.logistic.b_begin,
            b_inside=hmm.logistic.b_inside,
            first_paragraph_bias=hmm.logistic.first_paragraph_bias,
        )

    def brain(self) -> CompanyBrainMindMap:
        return self._brain


registry = AlgorithmRegistry()
store = Store()
store.reset()
init_db()
request_role: ContextVar[Role | None] = ContextVar("request_role", default=None)


def set_request_role(role: Role) -> Token[Role | None]:
    return request_role.set(role)


def reset_request_role(token: Token[Role | None]) -> None:
    request_role.reset(token)


def active_role() -> Role:
    return request_role.get() or store.current_role


def role_at_least(role: Role) -> bool:
    return ROLE_RANK[active_role()] >= ROLE_RANK[role]


def hash_api_key(key: str) -> str:
    return hashlib.sha256(key.encode("utf-8")).hexdigest()


def create_api_key(name: str, role: Role) -> dict[str, Any]:
    clean_name = name.strip()
    if not clean_name:
        raise ValueError("API key name is required.")
    token = f"lou_{secrets.token_urlsafe(32)}"
    record = ApiKeyRecord(
        id=f"key-{uuid.uuid4().hex[:10]}",
        name=clean_name,
        key_hash=hash_api_key(token),
        key_prefix=f"{token[:8]}...",
        role=role.value,
    )
    with Session(engine) as session:
        session.add(record)
        session.commit()
        session.refresh(record)
    return {
        "id": record.id,
        "name": record.name,
        "role": record.role,
        "key_prefix": record.key_prefix,
        "created_at": record.created_at,
        "key": token,
    }


def list_api_keys() -> list[dict[str, Any]]:
    with Session(engine) as session:
        records = session.exec(select(ApiKeyRecord).where(ApiKeyRecord.revoked_at.is_(None))).all()
    return [
        {
            "id": record.id,
            "name": record.name,
            "role": record.role,
            "key_prefix": record.key_prefix,
            "created_at": record.created_at,
            "last_used_at": record.last_used_at,
        }
        for record in records
    ]


def revoke_api_key(key_id: str) -> bool:
    with Session(engine) as session:
        record = session.get(ApiKeyRecord, key_id)
        if record is None or record.revoked_at is not None:
            return False
        record.revoked_at = datetime.now(timezone.utc).isoformat()
        session.add(record)
        session.commit()
    return True


def role_for_api_key(token: str) -> Role | None:
    key_hash = hash_api_key(token)
    with Session(engine) as session:
        record = session.exec(
            select(ApiKeyRecord).where(ApiKeyRecord.key_hash == key_hash, ApiKeyRecord.revoked_at.is_(None))
        ).first()
        if record is None:
            return None
        record.last_used_at = datetime.now(timezone.utc).isoformat()
        session.add(record)
        session.commit()
        return Role(record.role)


def summarize_playbook(playbook: Playbook) -> dict[str, Any]:
    return {
        "id": playbook.id,
        "slug": playbook.slug,
        "name": playbook.name,
        "category": playbook.category,
        "description": playbook.description,
        "owner": playbook.owner,
        "version": playbook.version,
        "position_count": len(playbook.positions),
    }


def persist_state() -> None:
    payload = {
        "playbooks": [playbook.model_dump() for playbook in store.playbooks.values()],
        "contracts": [contract.model_dump() for contract in store.contracts.values()],
        "proposals": [proposal.model_dump() for proposal in store.proposals.values()],
        "commits": [commit.model_dump() for commit in store.commits.values()],
    }
    write_snapshot("lou-demo-state", json.dumps(payload))


def _normalize(name: str) -> str:
    return "".join(ch.lower() for ch in name if ch.isalnum())


def _column_value(values: dict[str, str], name: str) -> str:
    normalised = _normalize(name)
    for key, value in values.items():
        if _normalize(key) == normalised:
            return value
    return ""


def _first_column_value(values: dict[str, str], names: list[str]) -> str:
    for name in names:
        value = _column_value(values, name)
        if value:
            return value
    return ""


def _infer_keywords(values: dict[str, str]) -> list[str]:
    text = " ".join(values.values())
    words = re.findall(r"[A-Za-z][A-Za-z-]{4,}", text)
    stop = {"position", "fallback", "preferred", "party", "legal", "terms", "clause", "agreement"}
    seen: list[str] = []
    for word in words:
        lowered = word.lower()
        if lowered in stop or lowered in seen:
            continue
        seen.append(lowered)
        if len(seen) == 8:
            break
    return seen


def sync_position_from_columns(position: PlaybookPosition) -> None:
    position.topic = _column_value(position.columns, "Topic") or position.topic
    position.preferred_position = _column_value(position.columns, "Preferred Position") or position.preferred_position
    position.fallback_position = _first_column_value(
        position.columns, ["Fallback Position", "Fallback 1", "Fallback"]
    ) or position.fallback_position
    position.risk = _first_column_value(position.columns, ["Risk", "Red Line", "Deal Breaker"]) or position.risk or "Medium"
    keywords = _first_column_value(position.columns, ["Keywords"])
    if keywords:
        position.keywords = [item.strip().lower() for item in keywords.split(";") if item.strip()]
    elif not position.keywords:
        position.keywords = _infer_keywords(position.columns)


def update_playbook_position(playbook_id: str, position_id: str, columns: dict[str, str]) -> PlaybookPosition:
    playbook = store.playbooks[playbook_id]
    position = next((item for item in playbook.positions if item.id == position_id), None)
    if position is None:
        raise KeyError(position_id)

    allowed_columns = set(playbook.columns)
    position.columns = {column: str(columns.get(column, position.columns.get(column, ""))) for column in playbook.columns}
    for column, value in columns.items():
        if column in allowed_columns:
            continue
        position.columns[column] = str(value)
        playbook.columns.append(column)

    sync_position_from_columns(position)
    playbook.version += 1
    registry.invalidate(playbook_id)
    persist_state()
    return position


def split_clauses(text: str) -> list[str]:
    raw = re.split(r"(?<=[.;:!?])\s+|\n+", text.strip())
    clauses = [item.strip() for item in raw if len(item.strip()) > 12]
    return clauses or [text.strip()]


def split_paragraphs(text: str) -> list[str]:
    return [chunk.strip() for chunk in re.split(r"\n{2,}", text) if chunk.strip()]


def sectionize_text(text: str) -> list[dict[str, Any]]:
    detector = registry.section_detector()
    paragraphs = split_paragraphs(text)
    if not paragraphs:
        return []
    sections = detector.segment(paragraphs)

    payload: list[dict[str, Any]] = []
    cursor = 0
    for section in sections:
        start = text.find(section.text, cursor)
        if start < 0:
            start = cursor
        end = start + len(section.text)
        payload.append(
            {
                "id": f"section-{section.index}",
                "title": section.title or f"Section {section.index}",
                "text": section.text,
                "start": start,
                "end": end,
            }
        )
        cursor = end
    return payload


def location_for(text: str, excerpt: str, sections: list[dict[str, Any]]) -> str:
    pos = text.find(excerpt)
    if pos < 0:
        return "Unlocated clause"
    for section in sections:
        if section["start"] <= pos < section["end"]:
            return section["title"]
    return f"Section near char {pos}"


def build_highlights(contract: Contract) -> list[dict[str, Any]]:
    highlights = []
    search_from = 0
    for finding in contract.findings:
        start = contract.text.find(finding.excerpt, search_from)
        if start < 0:
            start = contract.text.find(finding.excerpt)
        if start < 0:
            continue
        end = start + len(finding.excerpt)
        finding.start = start
        finding.end = end
        search_from = end
        highlights.append(
            {
                "id": f"hl-{finding.id}",
                "finding_id": finding.id,
                "status": finding.status.value,
                "topic": finding.topic,
                "start": start,
                "end": end,
                "excerpt": finding.excerpt,
                "risk": finding.risk,
            }
        )
    return highlights


def create_proposal(
    playbook_id: str,
    topic: str,
    source: str,
    proposed_text: str,
    rationale: str,
    voice_match_scores: dict | None = None,
    voice_session_id: str | None = None,
) -> Proposal:
    proposal = Proposal(
        id=f"prop-{uuid.uuid4().hex[:10]}",
        playbook_id=playbook_id,
        topic=topic,
        source=source,
        proposed_text=proposed_text,
        rationale=rationale,
        created_by_role=active_role(),
        voice_match_scores=voice_match_scores,
        voice_session_id=voice_session_id,
        created_at=datetime.now(timezone.utc).isoformat(),
    )
    store.proposals[proposal.id] = proposal
    return proposal


def suggested_proposal(
    playbook_id: str,
    topic: str,
    source: str,
    proposed_text: str,
    rationale: str,
    voice_match_scores: dict | None = None,
    voice_session_id: str | None = None,
) -> Proposal:
    return Proposal(
        id=f"suggestion-{uuid.uuid4().hex[:10]}",
        playbook_id=playbook_id,
        topic=topic,
        source=source,
        proposed_text=proposed_text,
        rationale=rationale,
        created_by_role=active_role(),
        voice_match_scores=voice_match_scores,
        voice_session_id=voice_session_id,
        created_at=datetime.now(timezone.utc).isoformat(),
    )


def _find_position(playbook: Playbook, position_id: str) -> PlaybookPosition | None:
    return next((item for item in playbook.positions if item.id == position_id), None)


def analyze_contract(
    playbook_id: str,
    name: str,
    text: str,
    sections: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    playbook = store.playbooks[playbook_id]
    matcher = registry.get_matcher(playbook)
    risk_scorer = registry.risk_scorer()

    contract = Contract(
        id=f"con-{uuid.uuid4().hex[:10]}",
        playbook_id=playbook_id,
        name=name,
        text=text,
        sections=sections or sectionize_text(text),
    )
    proposed_updates: list[Proposal] = []
    observations: list[RiskObservation] = []

    for index, clause in enumerate(split_clauses(text), start=1):
        match = matcher.best_match(clause)
        if match is not None:
            position = _find_position(playbook, match.position_id)
            finding = ContractFinding(
                id=f"find-{uuid.uuid4().hex[:10]}",
                contract_id=contract.id,
                playbook_position_id=match.position_id,
                topic=match.topic,
                excerpt=clause,
                status=FindingStatus.MAPPED,
                risk=position.risk if position else "Medium",
                location=location_for(text, clause, contract.sections),
                recommendation=position.preferred_position if position else "",
                match_score=match.score,
                match_method=match.method,
            )
            observations.append(RiskObservation(risk_label=finding.risk, weight=match.score))
        else:
            topic = infer_topic(clause)
            risk_label = "High" if re.search(r"non-solicit|penalty|perpetual|unlimited", clause, re.IGNORECASE) else "Medium"
            finding = ContractFinding(
                id=f"find-{uuid.uuid4().hex[:10]}",
                contract_id=contract.id,
                playbook_position_id=None,
                topic=topic,
                excerpt=clause,
                status=FindingStatus.UNMAPPED,
                risk=risk_label,
                location=f"Unmapped clause {index}",
                recommendation="No playbook position matched. Propose guidance before treating this clause as acceptable.",
                match_score=0.0,
                match_method="unmapped",
            )
            observations.append(RiskObservation(risk_label=risk_label, weight=0.5))
            proposed_updates.append(
                suggested_proposal(
                    playbook_id=playbook_id,
                    topic=topic,
                    source="contract",
                    proposed_text=f"Create guidance for {topic}: {clause[:200]}",
                    rationale="Contract analysis found language outside the current playbook coverage.",
                )
            )
        contract.findings.append(finding)

    contract.highlights = build_highlights(contract)
    posterior = risk_scorer.score(observations)
    contract.risk_posterior = posterior.to_dict()
    store.contracts[contract.id] = contract
    persist_state()
    return {
        "contract": contract,
        "sections": contract.sections,
        "findings": contract.findings,
        "highlights": contract.highlights,
        "risk_posterior": contract.risk_posterior,
        "proposed_updates": proposed_updates,
        "review_suggestions": review_suggestions(contract),
    }


def review_suggestions(contract: Contract) -> list[dict[str, str]]:
    suggestions = []
    for finding in contract.findings:
        if finding.status == FindingStatus.MAPPED:
            suggestions.append(
                {
                    "finding_id": finding.id,
                    "title": f"{finding.topic} is covered",
                    "body": f"Confirm the clause tracks the preferred position (match {finding.match_score:.2f}).",
                }
            )
        else:
            suggestions.append(
                {
                    "finding_id": finding.id,
                    "title": f"Add guidance for {finding.topic}",
                    "body": "Route this clause to the review queue before treating it as approved.",
                }
            )
    return suggestions


def _reject_zip_bombs(filename: str, content: bytes, max_ratio: int) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            total_uncompressed = 0
            total_compressed = 0
            for info in zf.infolist():
                total_uncompressed += info.file_size
                total_compressed += info.compress_size
            if total_compressed > 0:
                ratio = total_uncompressed / total_compressed
                if ratio > max_ratio:
                    raise ValueError(
                        f"{filename} is suspiciously compressed (ratio {ratio:.0f}x > {max_ratio}x cap)."
                    )
    except zipfile.BadZipFile as error:
        raise ValueError(f"{filename} is not a valid DOCX archive.") from error


def extract_pdf_document(filename: str, content: bytes) -> tuple[str, list[dict[str, Any]]]:
    from .limits import MAX_PDF_PAGES  # local import avoids cycle at module load

    reader = PdfReader(io.BytesIO(content))
    if len(reader.pages) > MAX_PDF_PAGES:
        raise ValueError(
            f"{filename} has {len(reader.pages)} pages; the cap is {MAX_PDF_PAGES}."
        )
    sections = []
    parts = []
    cursor = 0
    for index, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if not page_text:
            continue
        if parts:
            parts.append("\n\n")
            cursor += 2
        start = cursor
        parts.append(page_text)
        cursor += len(page_text)
        sections.append({"id": f"page-{index}", "title": f"Page {index}", "text": page_text, "start": start, "end": cursor})
    text = "".join(parts).strip()
    if not text:
        raise ValueError(f"{filename} has no extractable text. Scanned-image OCR is not supported in this build.")
    return text, sections


def extract_docx_document(filename: str, content: bytes) -> tuple[str, list[dict[str, Any]]]:
    from .limits import DOCX_MAX_RATIO, MAX_DOCX_PARAGRAPHS  # local import avoids cycle

    _reject_zip_bombs(filename, content, DOCX_MAX_RATIO)
    document = Document(io.BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    if len(paragraphs) > MAX_DOCX_PARAGRAPHS:
        raise ValueError(
            f"{filename} has {len(paragraphs)} paragraphs; the cap is {MAX_DOCX_PARAGRAPHS}."
        )
    sections = []
    parts = []
    cursor = 0
    for index, paragraph in enumerate(paragraphs, start=1):
        if parts:
            parts.append("\n\n")
            cursor += 2
        start = cursor
        parts.append(paragraph)
        cursor += len(paragraph)
        sections.append({"id": f"section-{index}", "title": f"Section {index}", "text": paragraph, "start": start, "end": cursor})
    text = "".join(parts).strip()
    if not text:
        raise ValueError(f"{filename} has no extractable text.")
    return text, sections


def analyze_uploaded_contract(playbook_id: str, filename: str, content: bytes) -> dict[str, Any]:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        text, sections = extract_pdf_document(filename, content)
    elif lower.endswith(".docx"):
        text, sections = extract_docx_document(filename, content)
    else:
        raise ValueError("Upload a text-based PDF or DOCX contract.")
    return analyze_contract(playbook_id, filename, text, sections)


def build_contract_review_artifact(playbook_id: str, filename: str, content: bytes) -> StreamingResponse:
    result = analyze_uploaded_contract(playbook_id, filename, content)
    warnings: list[str] = []
    annotated_name, annotated_content = annotate_contract_document(filename, content, result["findings"], warnings)
    review_payload = {
        "contract": result["contract"],
        "findings": result["findings"],
        "highlights": result["highlights"],
        "risk_posterior": result["risk_posterior"],
        "warnings": warnings,
    }
    output = io.BytesIO()
    with zipfile.ZipFile(output, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(annotated_name, annotated_content)
        archive.writestr("review.json", json.dumps(jsonable_encoder(review_payload), indent=2))
    output.seek(0)
    base = Path(filename).stem or "contract"
    return StreamingResponse(
        output,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{base}-lou-review.zip"'},
    )


def annotate_contract_document(
    filename: str,
    content: bytes,
    findings: list[ContractFinding],
    warnings: list[str],
) -> tuple[str, bytes]:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return f"annotated-{Path(filename).stem}.pdf", annotate_pdf_contract(filename, content, findings, warnings)
    if lower.endswith(".docx"):
        return f"annotated-{Path(filename).stem}.docx", annotate_docx_contract(filename, content, findings, warnings)
    raise ValueError("Upload a text-based PDF or DOCX contract.")


def annotate_pdf_contract(filename: str, content: bytes, findings: list[ContractFinding], warnings: list[str]) -> bytes:
    document = fitz.open(stream=content, filetype="pdf")
    for finding in findings:
        excerpt = finding.excerpt.strip()
        if not excerpt:
            continue
        matched = False
        for page in document:
            for rect in page.search_for(excerpt):
                annotation = page.add_highlight_annot(rect)
                annotation.set_info(title="Lou", content=f"{finding.topic}: {finding.recommendation}")
                annotation.update()
                matched = True
        if not matched:
            warnings.append(f"Excerpt not present in {filename}: {excerpt[:120]}")
    output = io.BytesIO()
    document.save(output)
    document.close()
    return output.getvalue()


def annotate_docx_contract(filename: str, content: bytes, findings: list[ContractFinding], warnings: list[str]) -> bytes:
    document = Document(io.BytesIO(content))
    for finding in findings:
        excerpt = finding.excerpt.strip()
        if not excerpt:
            continue
        matched = False
        for paragraph in document.paragraphs:
            if excerpt not in paragraph.text:
                continue
            matched = True
            run_matched = False
            for run in paragraph.runs:
                if excerpt in run.text:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    run_matched = True
            if not run_matched:
                for run in paragraph.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                warnings.append(
                    f"Exact run boundaries were not preserved for excerpt in {filename}: {excerpt[:120]}"
                )
        if not matched:
            warnings.append(f"Excerpt not present in {filename}: {excerpt[:120]}")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def infer_topic(text: str) -> str:
    lower = text.lower()
    if "non-solicit" in lower or "employee" in lower:
        return "Expanded Non-Solicit"
    if "residual" in lower:
        return "Residual Knowledge"
    if "archival" in lower:
        return "Archival Retention"
    words = [word.title() for word in re.findall(r"[A-Za-z][A-Za-z-]{4,}", text)[:3]]
    return " ".join(words) or "Unmapped Clause"


def normalize_language(language: str) -> str:
    languages = set(settings.VOICE_LANGUAGES)
    return language if language in languages else "en"


def voice_session(playbook_id: str | None, language: str = "en") -> dict[str, Any]:
    api_key_present = bool(current_slng_key(settings.SLNG_API_KEY))
    selected_language = normalize_language(language)
    return {
        "provider": "SLNG",
        "mode": "live" if api_key_present else "transcript-fallback",
        "playbook_id": playbook_id,
        "language": selected_language,
        "supported_languages": sorted(set(settings.VOICE_LANGUAGES)),
        "stt": {
            "model": settings.SLNG_STT_MODEL,
            "http_url": _slng_stt_http_url(),
            "websocket_url": _slng_stt_ws_url(),
            "request": {
                "language": selected_language,
                "diarize": True,
                "punctuate": True,
                "smart_format": True,
                "utterances": True,
                "keywords": list(settings.SLNG_STT_KEYWORDS),
            },
        },
        "tts": {
            "model": settings.SLNG_TTS_MODEL,
            "http_url": _slng_tts_http_url(),
            "websocket_url": _slng_tts_ws_url(),
            "request": {
                "speaker": settings.SLNG_TTS_SPEAKER,
                "text": settings.SLNG_TTS_GREETING,
                "audioFormat": "audio/wav",
                "sample_rate": settings.SLNG_TTS_SAMPLE_RATE,
                "speed": settings.SLNG_TTS_SPEED,
            },
        },
        "auth_note": "Browser WebSocket clients cannot set Authorization headers; proxy through the backend or pass a short-lived token as a query parameter.",
        "api_key_present": api_key_present,
    }


def extract_slng_speaker_segments(payload: dict[str, Any]) -> list[dict[str, Any]]:
    utterances = payload.get("utterances")
    if not isinstance(utterances, list):
        results = payload.get("results")
        utterances = results.get("utterances") if isinstance(results, dict) else None
    if isinstance(utterances, list):
        segments = []
        for index, utterance in enumerate(utterances, start=1):
            if not isinstance(utterance, dict):
                continue
            text = str(utterance.get("transcript") or utterance.get("text") or "").strip()
            if not text:
                continue
            speaker = utterance.get("speaker")
            segments.append(
                {
                    "speaker": f"Speaker {int(speaker) + 1}" if isinstance(speaker, int) else f"Speaker {speaker or index}",
                    "text": text,
                    "start": utterance.get("start"),
                    "end": utterance.get("end"),
                }
            )
        if segments:
            return segments

    words: list[dict[str, Any]] = []
    results = payload.get("results")
    if isinstance(results, dict):
        channels = results.get("channels")
        if isinstance(channels, list):
            for channel in channels:
                if not isinstance(channel, dict):
                    continue
                alternatives = channel.get("alternatives", [])
                if not isinstance(alternatives, list):
                    continue
                for alternative in alternatives:
                    if isinstance(alternative, dict) and isinstance(alternative.get("words"), list):
                        words.extend(word for word in alternative["words"] if isinstance(word, dict))

    segments: list[dict[str, Any]] = []
    current_speaker: Any = None
    current_words: list[str] = []
    current_start: Any = None
    current_end: Any = None
    for word in words:
        speaker = word.get("speaker", 0)
        token = str(word.get("punctuated_word") or word.get("word") or "").strip()
        if not token:
            continue
        if current_words and speaker != current_speaker:
            speaker_label = f"Speaker {int(current_speaker) + 1}" if isinstance(current_speaker, int) else f"Speaker {current_speaker}"
            segments.append({"speaker": speaker_label, "text": " ".join(current_words), "start": current_start, "end": current_end})
            current_words = []
            current_start = None
        if not current_words:
            current_speaker = speaker
            current_start = word.get("start")
        current_words.append(token)
        current_end = word.get("end")
    if current_words:
        speaker_label = f"Speaker {int(current_speaker) + 1}" if isinstance(current_speaker, int) else f"Speaker {current_speaker}"
        segments.append({"speaker": speaker_label, "text": " ".join(current_words), "start": current_start, "end": current_end})
    return segments


def extract_slng_transcript(payload: dict[str, Any]) -> str:
    speaker_segments = extract_slng_speaker_segments(payload)
    if speaker_segments:
        return "\n".join(f"{segment['speaker']}: {segment['text']}" for segment in speaker_segments)

    for key in ("text", "transcript"):
        value = payload.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip()

    results = payload.get("results")
    if isinstance(results, dict):
        channels = results.get("channels")
        if isinstance(channels, list):
            alternatives = []
            for channel in channels:
                if isinstance(channel, dict):
                    alternatives.extend(channel.get("alternatives", []))
            for alternative in alternatives:
                if isinstance(alternative, dict):
                    transcript = alternative.get("transcript")
                    if isinstance(transcript, str) and transcript.strip():
                        return transcript.strip()

    raise ValueError("SLNG transcription response did not include transcript text.")


def transcribe_audio_with_slng(
    audio: bytes,
    filename: str,
    content_type: str | None,
    language: str = "en",
) -> dict[str, Any]:
    slng_key = current_slng_key(settings.SLNG_API_KEY)
    if not slng_key:
        raise ValueError("Set LOU_SLNG_API_KEY or SLNG_API_KEY to enable recorded-audio transcription.")
    if not audio:
        raise ValueError("Recorded audio was empty.")

    selected_language = normalize_language(language)
    upload_name = filename or "lou-recording.webm"
    upload_type = content_type or "application/octet-stream"

    try:
        response = httpx.post(
            _slng_stt_http_url(),
            headers={"Authorization": f"Bearer {slng_key}"},
            files={"audio": (upload_name, audio, upload_type)},
            data={
                "language": selected_language,
                "diarize": "true",
                "punctuate": "true",
                "smart_format": "true",
                "utterances": "true",
            },
            timeout=60,
        )
        response.raise_for_status()
    except httpx.HTTPStatusError as error:
        detail = error.response.text[:500]
        raise ValueError(f"SLNG transcription failed with {error.response.status_code}: {detail}") from error
    except httpx.HTTPError as error:
        raise ValueError(f"SLNG transcription request failed: {error}") from error

    payload = response.json()
    transcript = extract_slng_transcript(payload)
    speaker_segments = extract_slng_speaker_segments(payload)
    return {
        "provider": "SLNG",
        "mode": "slng-audio",
        "language": selected_language,
        "transcript": transcript,
        "speaker_segments": speaker_segments,
        "raw": payload,
    }


def transcribe_audio_to_updates(
    playbook_id: str,
    audio: bytes,
    filename: str,
    content_type: str | None,
    language: str = "en",
) -> dict[str, Any]:
    # Graceful degradation: when no SLNG key is configured, return a 200 with an
    # empty transcript and a clear message instead of bubbling a 502 to the UI.
    # The Voice page already supports pasted-transcript fallback for this case.
    if not current_slng_key(settings.SLNG_API_KEY):
        return {
            "mode": "transcript-fallback",
            "provider": "none",
            "language": normalize_language(language),
            "transcript": "",
            "speaker_segments": [],
            "proposed_updates": [],
            "message": (
                "Live audio transcription is disabled because no SLNG key is configured."
                " Paste the transcript below to create proposals."
            ),
        }
    transcription = transcribe_audio_with_slng(audio, filename, content_type, language)
    result = transcript_to_updates(playbook_id, transcription["transcript"], transcription["language"])
    result["mode"] = transcription["mode"]
    result["provider"] = transcription["provider"]
    result["transcript"] = transcription["transcript"]
    result["speaker_segments"] = transcription["speaker_segments"]
    return result


def contract_from_voice_notes(playbook_id: str, transcript: str, language: str = "en") -> dict[str, Any]:
    if playbook_id not in store.playbooks:
        raise ValueError("Playbook not found.")
    notes = transcript.strip()
    if not notes:
        raise ValueError("Capture or paste speaker-separated notes before generating a contract.")

    playbook = store.playbooks[playbook_id]
    generated = draft_contract_from_notes(notes, playbook.name, playbook.category)
    analysis = analyze_contract(playbook_id, generated["title"], generated["contract_text"])
    analysis["mode"] = "openai-contract-from-notes"
    analysis["language"] = normalize_language(language)
    analysis["generated_contract"] = generated
    return analysis


def transcript_to_updates(playbook_id: str, transcript: str, language: str = "en") -> dict[str, Any]:
    playbook = store.playbooks[playbook_id]
    voice_matcher = registry.voice_matcher(playbook)
    candidates = [(position.id, position.topic, position.preferred_position) for position in playbook.positions]
    sentences = split_clauses(transcript) or [transcript.strip()]
    proposals: list[Proposal] = []

    for sentence in sentences:
        if not sentence:
            continue
        matches = voice_matcher.match(sentence, candidates, top_k=3)
        match_scores = [
            {
                "position_id": match.position_id,
                "topic": match.topic,
                "score": match.score,
                "jaro": match.jaro,
                "tfidf": match.tfidf,
                "edit": match.edit,
            }
            for match in matches
        ]
        topic = matches[0].topic if matches else infer_topic(sentence)
        proposals.append(
            suggested_proposal(
                playbook_id=playbook_id,
                topic=topic,
                source="voice",
                proposed_text=f"Discussion note: {sentence[:220]}",
                rationale="Captured from listening-mode transcript.",
                voice_match_scores={"matches": match_scores},
            )
        )

    if not proposals:
        proposals.append(
            suggested_proposal(
                playbook_id=playbook_id,
                topic="Meeting Follow-up",
                source="voice",
                proposed_text=f"Review transcript for possible playbook impact: {transcript[:220]}",
                rationale="Listening mode captured a legal discussion without a direct playbook match.",
            )
        )

    return {
        "mode": "transcript-fallback",
        "language": normalize_language(language),
        "proposed_updates": proposals,
    }


def approve_proposal(proposal_id: str, author_role: Role, edited_text: str | None = None) -> dict[str, Any]:
    proposal = store.proposals[proposal_id]
    if edited_text:
        proposal.proposed_text = edited_text
    proposal.status = ProposalStatus.APPROVED
    commit = Commit(
        id=f"commit-{uuid.uuid4().hex[:8]}",
        proposal_id=proposal.id,
        playbook_id=proposal.playbook_id,
        message=f"Approved {proposal.topic}",
        author_role=author_role,
        created_at=datetime.now(timezone.utc).isoformat(),
    )
    store.commits[commit.id] = commit
    playbook = store.playbooks[proposal.playbook_id]
    playbook.version += 1
    registry.invalidate(proposal.playbook_id)
    persist_state()
    return {"proposal": proposal, "commit": commit}


def reject_proposal(proposal_id: str, reason: str | None) -> Proposal:
    proposal = store.proposals[proposal_id]
    proposal.status = ProposalStatus.REJECTED
    if reason:
        proposal.rationale = f"{proposal.rationale} Rejected: {reason}"
    persist_state()
    return proposal


def graph_for_playbook(playbook: Playbook) -> dict[str, Any]:
    """Render a single playbook as a .mm-style mini-brain tree."""
    snapshot = registry.brain().from_playbook(playbook)
    nodes = list(snapshot.nodes)
    edges = [{**edge, "kind": "hierarchy"} for edge in snapshot.edges]
    topic_ids = {position.topic.lower(): f"topic-{position.id}" for position in playbook.positions}

    for source, target, strength in _topic_relation_edges(playbook):
        edges.append(
            {
                "source": f"topic-{source.id}",
                "target": f"topic-{target.id}",
                "kind": "topic_relation",
                "label": "related",
                "strength": strength,
            }
        )

    for proposal in store.proposals.values():
        if proposal.playbook_id != playbook.id:
            continue
        parent = topic_ids.get(proposal.topic.lower(), playbook.id)
        nodes.append(
            {
                "id": proposal.id,
                "name": proposal.topic[:24],
                "label": proposal.topic[:24],
                "kind": "proposal",
                "type": "leaf",
                "summary": proposal.status.value,
                "metadata": {"playbook_id": playbook.id, "proposal_id": proposal.id},
                "depth": 3,
            }
        )
        edges.append({"source": parent, "target": proposal.id, "kind": "hierarchy", "label": "proposed"})

    for commit in store.commits.values():
        if commit.playbook_id != playbook.id:
            continue
        proposal = store.proposals.get(commit.proposal_id)
        nodes.append(
            {
                "id": commit.id,
                "name": "Published",
                "label": "Published",
                "kind": "commit",
                "type": "leaf",
                "summary": commit.message,
                "metadata": {"playbook_id": playbook.id, "proposal_id": commit.proposal_id},
                "depth": 4,
            }
        )
        edges.append(
            {
                "source": proposal.id if proposal else playbook.id,
                "target": commit.id,
                "kind": "hierarchy",
                "label": "approved",
            }
        )

    depth_counts: dict[int, int] = {}
    for node in nodes:
        depth = int(node.get("depth", 0))
        order = depth_counts.get(depth, 0)
        depth_counts[depth] = order + 1
        node.setdefault("x", round(80 + depth * 180, 2))
        node.setdefault("y", round(44 + order * 36, 2))

    metrics = {
        **snapshot.metrics,
        "node_count": len(nodes),
        "edge_count": len(edges),
        "relation_count": sum(1 for edge in edges if edge.get("kind") == "topic_relation"),
    }
    return {"nodes": nodes, "edges": edges, "tree": snapshot.tree, "metrics": metrics}


def _short_summary(topic: str) -> str:
    words = re.findall(r"[A-Za-z][A-Za-z-]+", topic)
    stop = {"of", "and", "or", "the", "for", "to"}
    selected = [word for word in words if word.lower() not in stop][:3]
    return " ".join(selected) or topic[:32]


def _topic_relation_edges(playbook: Playbook) -> list[tuple[PlaybookPosition, PlaybookPosition, float]]:
    relations = []
    positions = playbook.positions
    for index, source in enumerate(positions):
        source_terms = set(source.keywords) | {
            word.lower() for word in re.findall(r"[A-Za-z][A-Za-z-]{4,}", source.topic)
        }
        for target in positions[index + 1 :]:
            target_terms = set(target.keywords) | {
                word.lower() for word in re.findall(r"[A-Za-z][A-Za-z-]{4,}", target.topic)
            }
            overlap = source_terms & target_terms
            if overlap:
                denominator = max(min(len(source_terms), len(target_terms)), 1)
                strength = round(min(1.0, len(overlap) / denominator), 2)
                relations.append((source, target, strength))
    if not relations and len(positions) >= 2:
        relations.append((positions[0], positions[1], 0.2))
    return sorted(relations, key=lambda item: item[2], reverse=True)[:24]


def company_brain() -> dict[str, Any]:
    playbooks = list(store.playbooks.values())
    if not playbooks:
        return {
            "nodes": [],
            "edges": [],
            "tree": {"id": "company-brain", "name": "Company Brain", "type": "branch", "children": []},
            "metrics": {"node_count": 0, "edge_count": 0, "branch_count": 0, "relation_count": 0},
        }
    snapshot = registry.brain().from_playbooks(playbooks)
    return {"nodes": snapshot.nodes, "edges": snapshot.edges, "tree": snapshot.tree, "metrics": snapshot.metrics}


def export_json() -> StreamingResponse:
    payload = {
        "playbooks": [playbook.model_dump() for playbook in store.playbooks.values()],
        "proposals": [proposal.model_dump() for proposal in store.proposals.values()],
        "commits": [commit.model_dump() for commit in store.commits.values()],
    }
    content = json.dumps(payload, indent=2)
    return StreamingResponse(io.BytesIO(content.encode("utf-8")), media_type="application/json")


def export_xlsx() -> StreamingResponse:
    wb = Workbook()
    ws = wb.active
    ws.title = "Playbook Positions"
    for playbook in store.playbooks.values():
        columns = ["Playbook", "Version", *playbook.columns]
        ws.append(columns)
        for position in playbook.positions:
            ws.append([playbook.name, playbook.version, *[position.columns.get(column, "") for column in playbook.columns]])
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=lou-playbooks.xlsx"},
    )


def export_png_placeholder() -> StreamingResponse:
    data = bytes.fromhex(
        "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
        "0000000a49444154789c6360000002000100ffff03000006000557bfab0000000049454e44ae426082"
    )
    return StreamingResponse(io.BytesIO(data), media_type="image/png")


async def parse_lou_command(command: str, playbook_id: str | None) -> dict[str, Any]:
    try:
        ai_result = parse_command_with_openai(command)
    except Exception:
        ai_result = None
    if ai_result:
        intent = str(ai_result.get("intent") or "note")
        return {"intent": intent, "playbook_id": playbook_id, "message": command_message(intent)}

    lower = command.lower()
    if "approve" in lower:
        return {"intent": "approve", "message": command_message("approve")}
    if "export" in lower:
        return {"intent": "export", "message": command_message("export")}
    if "analyze" in lower or "review" in lower:
        return {"intent": "analyze_contract", "message": command_message("analyze_contract")}
    return {"intent": "note", "playbook_id": playbook_id, "message": command_message("note")}


def command_message(intent: str) -> str:
    messages = {
        "approve": "Open the review queue and approve the matching proposal.",
        "reject": "Open the review queue and reject the matching proposal.",
        "export": "Open exports for JSON, XLSX, or graph image.",
        "analyze_contract": "Upload a contract to start review.",
        "open_playbook": "Open the playbook workspace.",
        "note": "Saved as a workspace note.",
    }
    return messages.get(intent, messages["note"])
